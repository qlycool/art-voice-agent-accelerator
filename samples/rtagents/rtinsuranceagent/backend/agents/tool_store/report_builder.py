"""
Generate a polished FNOL report as a .docx file.
Call generate_claim_report(claim_dict) right after record_fnol success.
"""
from datetime import datetime
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.shared import Inches, Pt

OUT_DIR = Path("./fnol_reports")
OUT_DIR.mkdir(exist_ok=True)


def _h(doc, text, lvl=0):
    p = doc.add_heading(text, level=lvl)
    p.alignment = 0  # left


def _add_kv(doc, k, v):
    row = doc.add_paragraph()
    row.add_run(f"{k}: ").bold = True
    row.add_run(str(v))


def generate_claim_report(claim: Dict[str, Any]) -> Path:
    """
    claim – dict as stored in claims_db (includes claim_id, status, etc.)
    returns  Path to the .docx file.
    """
    doc = Document()
    doc.core_properties.title = f"FNOL Report – {claim['claim_id']}"

    # --- Title page ---------------------------------------------------
    _h(doc, "First Notice of Loss (FNOL) Report", 0)
    doc.add_paragraph(f"Claim ID: {claim['claim_id']}")
    doc.add_paragraph(f"Policy ID: {claim['policy_id']}")
    doc.add_paragraph(
        f"Generated: {datetime.utcnow().isoformat(timespec='seconds')} UTC"
    )
    doc.add_page_break()

    # --- Section: Claim Summary --------------------------------------
    _h(doc, "Claim Summary", 1)
    summary_fields = [
        ("Caller Name", claim["caller_name"]),
        ("Caller Role", claim["caller_role"]),
        ("Date Reported", claim["date_reported"]),
        ("Date of Loss", claim["date_of_loss"]),
        ("Time of Loss", claim.get("time_of_loss") or "N/A"),
        ("Collision", claim["collision"]),
        ("Comprehensive Loss", claim["comprehensive_loss"]),
        ("Bodily Injury", claim["bodily_injury"]),
        ("Property Damage", claim["property_damage"]),
        ("Status", claim["status"]),
    ]
    for k, v in summary_fields:
        _add_kv(doc, k, v)

    # --- Section: Loss Location --------------------------------------
    _h(doc, "Loss Location", 1)
    loc = claim["loss_location"]
    _add_kv(doc, "Street", loc["street"])
    _add_kv(doc, "City", loc["city"])
    _add_kv(doc, "State", loc["state"])
    _add_kv(doc, "ZIP", loc["zipcode"])
    if claim.get("location_description"):
        _add_kv(doc, "Description", claim["location_description"])

    # --- Section: Narrative ------------------------------------------
    _h(doc, "Narrative", 1)
    doc.add_paragraph(claim["narrative"])

    # --- Save ---------------------------------------------------------
    fname = OUT_DIR / f"{claim['claim_id']}.docx"
    doc.save(fname)
    return fname
